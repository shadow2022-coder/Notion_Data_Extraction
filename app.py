from flask import Flask, render_template_string, request, jsonify, send_file
import requests
import json
from datetime import date, timedelta, datetime
import io
from docx import Document

app = Flask(__name__)

# Note: For larger apps, it is recommended to move this to a separate templates/index.html file
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Notion Data Extractor</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }

        .container {
            max-width: 900px;
            margin: 0 auto;
            background: white;
            border-radius: 20px;
            padding: 40px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
        }

        .header {
            text-align: center;
            margin-bottom: 30px;
        }

        .header h1 {
            color: #333;
            font-size: 32px;
            margin-bottom: 10px;
        }

        .header p {
            color: #666;
        }

        .form-group {
            margin-bottom: 20px;
        }

        label {
            display: block;
            color: #333;
            font-weight: 600;
            margin-bottom: 8px;
            font-size: 14px;
        }

        input[type="text"],
        input[type="date"],
        input[type="number"],
        select {
            width: 100%;
            padding: 12px 15px;
            border: 2px solid #e0e0e0;
            border-radius: 10px;
            font-size: 14px;
            transition: border-color 0.3s;
        }

        input:focus, select:focus {
            outline: none;
            border-color: #667eea;
        }

        .help-text {
            font-size: 12px;
            color: #999;
            margin-top: 5px;
        }

        .date-range {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
        }

        .button {
            width: 100%;
            padding: 15px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 10px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.2s;
        }

        .button:hover {
            transform: translateY(-2px);
        }

        .button:disabled {
            background: #ccc;
            cursor: not-allowed;
            transform: none;
        }

        .dynamic-fields {
            display: none;
        }

        .dynamic-fields.show {
            display: block;
        }

        .status {
            padding: 12px;
            border-radius: 8px;
            margin: 20px 0;
            display: none;
        }

        .status.show {
            display: block;
        }

        .status.success {
            background: #d4edda;
            color: #155724;
            border: 1px solid #c3e6cb;
        }

        .status.error {
            background: #f8d7da;
            color: #721c24;
            border: 1px solid #f5c6cb;
        }

        .status.info {
            background: #d1ecf1;
            color: #0c5460;
            border: 1px solid #bee5eb;
        }

        .loading {
            text-align: center;
            padding: 20px;
            display: none;
        }

        .loading.show {
            display: block;
        }

        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 0 auto 10px;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        .results {
            margin-top: 30px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 10px;
            display: none;
        }

        .results.show {
            display: block;
        }

        .results h3 {
            color: #333;
            margin-bottom: 15px;
        }

        .results pre {
            background: white;
            padding: 15px;
            border-radius: 8px;
            overflow: auto;
            max-height: 400px;
            font-size: 13px;
        }

        .download-buttons {
            display: none;
            gap: 10px;
            margin-top: 15px;
        }

        .download-buttons.show {
            display: flex;
        }

        .download-btn {
            flex: 1;
            padding: 10px;
            background: #28a745;
            color: white;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            font-weight: 600;
        }

        .download-btn:hover {
            background: #218838;
        }

        .info-box {
            background: #e7f3ff;
            border: 1px solid #b3d9ff;
            border-radius: 8px;
            padding: 15px;
            margin-bottom: 20px;
        }

        .info-box h4 {
            color: #004085;
            margin-bottom: 8px;
        }

        .info-box ul {
            margin-left: 20px;
            color: #004085;
            font-size: 13px;
            line-height: 1.8;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Notion Data Extractor</h1>
            <p>Extract your Notion database content with ease</p>
        </div>

        <div class="info-box">
            <h4>Setup Instructions:</h4>
            <ul>
                <li>Get Integration Token from <a href="https://www.notion.so/my-integrations" target="_blank">Notion Integrations</a></li>
                <li>Enter your Database ID (32 characters with dashes)</li>
                <li>Share your database with the integration</li>
            </ul>
        </div>

        <form id="extractForm">
            <div class="form-group">
                <label for="token">Notion Integration Token *</label>
                <input type="text" id="token" placeholder="ntn_... or secret_..." required>
                <div class="help-text">Your Notion API token</div>
            </div>

            <div class="form-group">
                <label for="databaseId">Database ID *</label>
                <input type="text" id="databaseId" placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx" required>
                <div class="help-text">32-character database ID with dashes</div>
            </div>

            <div class="form-group">
                <label for="dateProperty">Date Property Name</label>
                <input type="text" id="dateProperty" value="Date" placeholder="Date">
                <div class="help-text">Column name for dates (case-sensitive)</div>
            </div>

            <div class="form-group">
                <label for="personProperty">Person Property Name (Optional)</label>
                <input type="text" id="personProperty" value="Assignee" placeholder="Assignee">
                <div class="help-text">Column name for assignees</div>
            </div>

            <div class="form-group">
                <label for="extractMode">Extract Mode *</label>
                <select id="extractMode" required>
                    <option value="today">Today Only</option>
                    <option value="specific_date">Specific Date</option>
                    <option value="date_range">Date Range</option>
                    <option value="last_n_days">Last N Days</option>
                    <option value="all">All Data</option>
                </select>
            </div>

            <div id="specificDateField" class="form-group dynamic-fields">
                <label for="specificDate">Select Date</label>
                <input type="date" id="specificDate">
            </div>

            <div id="dateRangeFields" class="form-group dynamic-fields">
                <label>Date Range</label>
                <div class="date-range">
                    <input type="date" id="startDate" placeholder="Start">
                    <input type="date" id="endDate" placeholder="End">
                </div>
            </div>

            <div id="lastNDaysField" class="form-group dynamic-fields">
                <label for="lastNDays">Number of Days</label>
                <input type="number" id="lastNDays" value="7" min="1" max="365">
            </div>

            <button type="submit" class="button" id="extractBtn">Extract Data</button>
        </form>

        <div class="loading" id="loading">
            <div class="spinner"></div>
            <p>Extracting data from Notion...</p>
        </div>

        <div class="status" id="status"></div>

        <div class="download-buttons" id="downloadButtons">
            <button class="download-btn" onclick="downloadFile('json')">Download JSON</button>
            <button class="download-btn" onclick="downloadFile('txt')">Download TXT</button>
            <button class="download-btn" onclick="downloadFile('docx')">Download Word</button>
        </div>

        <div class="results" id="results">
            <h3>Extraction Results</h3>
            <pre id="resultsContent"></pre>
        </div>
    </div>

    <script>
        let extractedData = null;

        document.getElementById('extractMode').addEventListener('change', function() {
            document.querySelectorAll('.dynamic-fields').forEach(f => f.classList.remove('show'));
            
            if (this.value === 'specific_date') {
                document.getElementById('specificDateField').classList.add('show');
            } else if (this.value === 'date_range') {
                document.getElementById('dateRangeFields').classList.add('show');
            } else if (this.value === 'last_n_days') {
                document.getElementById('lastNDaysField').classList.add('show');
            }
        });

        function showStatus(message, type) {
            const status = document.getElementById('status');
            status.textContent = message;
            status.className = 'status show ' + type;
        }

        function showLoading(show) {
            document.getElementById('loading').classList.toggle('show', show);
            document.getElementById('extractBtn').disabled = show;
        }

        document.getElementById('extractForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const formData = {
                token: document.getElementById('token').value.trim(),
                database_id: document.getElementById('databaseId').value.trim(),
                date_property: document.getElementById('dateProperty').value.trim() || 'Date',
                person_property: document.getElementById('personProperty').value.trim() || 'Assignee',
                extract_mode: document.getElementById('extractMode').value,
                specific_date: document.getElementById('specificDate').value,
                start_date: document.getElementById('startDate').value,
                end_date: document.getElementById('endDate').value,
                last_n_days: document.getElementById('lastNDays').value
            };

            showLoading(true);
            document.getElementById('results').classList.remove('show');
            document.getElementById('downloadButtons').classList.remove('show');

            try {
                const response = await fetch('/extract', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(formData)
                });

                const data = await response.json();

                if (!response.ok) {
                    throw new Error(data.error || 'Extraction failed');
                }

                extractedData = data.data;

                const resultText = data.data.map((item, i) => 
                    `${i + 1}. ${item.title}\\n   Date: ${item.date}\\n   Assignee: ${item.assignee}\\n   Content length: ${item.content.length} chars`
                ).join('\\n\\n');

                document.getElementById('resultsContent').textContent = resultText;
                document.getElementById('results').classList.add('show');
                document.getElementById('downloadButtons').classList.add('show');

                showStatus(`Successfully extracted ${data.data.length} page(s).`, 'success');

            } catch (error) {
                showStatus(`Error: ${error.message}`, 'error');
            } finally {
                showLoading(false);
            }
        });

        async function downloadFile(format) {
            if (!extractedData) return;

            try {
                const response = await fetch('/download/' + format, {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({data: extractedData})
                });

                if (!response.ok) {
                    const err = await response.json();
                    throw new Error(err.error || 'Download failed');
                }

                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `notion_export_${new Date().toISOString().split('T')[0]}.${format}`;
                a.click();
                window.URL.revokeObjectURL(url);
            } catch (error) {
                showStatus(`Download failed: ${error.message}`, 'error');
            }
        }
    </script>
</body>
</html>
'''

def get_date_filter(mode, date_property, **kwargs):
    """Build date filter based on mode"""
    today = date.today().isoformat()
    
    if mode == 'today':
        tomorrow = (date.today() + timedelta(days=1)).isoformat()
        return {
            "and": [
                {"property": date_property, "date": {"on_or_after": today}},
                {"property": date_property, "date": {"before": tomorrow}}
            ]
        }
    elif mode == 'specific_date':
        specific = kwargs.get('specific_date')
        if not specific:
            return None
        next_day = (datetime.strptime(specific, "%Y-%m-%d") + timedelta(days=1)).date().isoformat()
        return {
            "and": [
                {"property": date_property, "date": {"on_or_after": specific}},
                {"property": date_property, "date": {"before": next_day}}
            ]
        }
    elif mode == 'date_range':
        start = kwargs.get('start_date')
        end = kwargs.get('end_date')
        if not start or not end:
            return None
        return {
            "and": [
                {"property": date_property, "date": {"on_or_after": start}},
                {"property": date_property, "date": {"on_or_before": end}}
            ]
        }
    elif mode == 'last_n_days':
        try:
            n_days = int(kwargs.get('last_n_days', 7))
        except Exception:
            n_days = 7
        start = (date.today() - timedelta(days=n_days - 1)).isoformat()
        return {
            "and": [
                {"property": date_property, "date": {"on_or_after": start}},
                {"property": date_property, "date": {"on_or_before": today}}
            ]
        }
    return None

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/extract', methods=['POST'])
def extract():
    try:
        data = request.get_json(force=True)

        token = data.get('token')
        database_id = data.get('database_id')
        date_property = data.get('date_property') or 'Date'
        person_property = data.get('person_property') or 'Assignee'
        extract_mode = data.get('extract_mode') or 'all'

        if not token or not database_id:
            return jsonify({"error": "token and database_id are required"}), 400
        
        headers = {
            'Authorization': f'Bearer {token}',
            'Content-Type': 'application/json',
            'Notion-Version': '2022-06-28'
        }
        
        # Build filter
        date_filter = get_date_filter(
            extract_mode, 
            date_property,
            specific_date=data.get('specific_date'),
            start_date=data.get('start_date'),
            end_date=data.get('end_date'),
            last_n_days=data.get('last_n_days')
        )
        
        payload = {"filter": date_filter} if date_filter else {}
        
        # Query database
        response = requests.post(
            f'https://api.notion.com/v1/databases/{database_id}/query',
            headers=headers,
            json=payload,
            timeout=30
        )
        
        if not response.ok:
            # Attempt to return Notion's message if available
            try:
                err = response.json()
                return jsonify({"error": err.get('message', 'API request failed')}), response.status_code
            except Exception:
                return jsonify({"error": "API request failed"}), response.status_code
        
        pages = response.json().get('results', [])
        
        if not pages:
            return jsonify({"error": "No pages found matching criteria"}), 404
        
        # Process pages
        processed_data = []
        
        for page in pages:
            props = page.get('properties', {})
            
            # Get title
            title = 'Untitled'
            for prop_name in ['Name', 'Title', 'name', 'title']:
                if prop_name in props and props[prop_name].get('title'):
                    title_field = props[prop_name]['title']
                    if isinstance(title_field, list) and title_field:
                        title = title_field[0].get('plain_text', title)
                    break
            
            # Get date
            page_date = 'No date'
            if date_property in props and props[date_property].get('date'):
                page_date = props[date_property]['date'].get('start', 'No date')
            
            # Get assignee
            assignee = 'Unassigned'
            if person_property in props and props[person_property].get('people'):
                people = props[person_property]['people']
                if isinstance(people, list) and people:
                    assignee = people[0].get('name') or people[0].get('id') or 'Unknown'
            
            # Fetch blocks
            blocks_response = requests.get(
                f'https://api.notion.com/v1/blocks/{page.get("id")}/children',
                headers=headers,
                timeout=30
            )
            
            content = ''
            if blocks_response.ok:
                blocks = blocks_response.json().get('results', [])
                text_parts = []
                
                for block in blocks:
                    block_type = block.get('type')
                    if block_type and block.get(block_type, {}).get('rich_text'):
                        texts = [t.get('plain_text', '') for t in block[block_type]['rich_text']]
                        text_parts.append(' '.join(texts))
                    else:
                        # handle paragraph which may be stored under 'paragraph' with 'rich_text'
                        if block.get('paragraph') and block['paragraph'].get('rich_text'):
                            texts = [t.get('plain_text', '') for t in block['paragraph']['rich_text']]
                            text_parts.append(' '.join(texts))
                
                content = '\n'.join([p for p in text_parts if p])
            
            processed_data.append({
                'page_id': page.get('id'),
                'title': title,
                'date': page_date,
                'assignee': assignee,
                'content': content,
                'url': f"https://www.notion.so/{page.get('id', '').replace('-', '')}"
            })
        
        return jsonify({"data": processed_data})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/download/<format>', methods=['POST'])
def download(format):
    try:
        payload = request.get_json(force=True) or {}
        data = payload.get('data', [])

        if not data:
            return jsonify({"error": "No data provided"}), 400

        # JSON export
        if format == 'json':
            output = io.BytesIO()
            output.write(json.dumps(data, indent=2, ensure_ascii=False).encode('utf-8'))
            output.seek(0)
            return send_file(
                output,
                mimetype='application/json',
                as_attachment=True,
                download_name=f'notion_export_{date.today().isoformat()}.json'
            )

        # TXT export
        elif format == 'txt':
            text = f"NOTION EXPORT - {date.today().isoformat()}\n"
            text += "=" * 80 + "\n\n"
            
            for item in data:
                text += "\n" + "=" * 80 + "\n"
                text += f"TITLE: {item.get('title', 'Untitled')}\n"
                text += f"DATE: {item.get('date', 'No date')}\n"
                text += f"BY: {item.get('assignee', 'Unassigned')}\n"
                text += f"URL: {item.get('url', '')}\n"
                text += "=" * 80 + "\n\n"
                text += item.get('content', '') + "\n\n"
            
            output = io.BytesIO()
            output.write(text.encode('utf-8'))
            output.seek(0)
            return send_file(
                output,
                mimetype='text/plain',
                as_attachment=True,
                download_name=f'notion_export_{date.today().isoformat()}.txt'
            )

        # DOCX export
        elif format == 'docx':
            doc = Document()
            doc.add_heading('Notion Export', level=1)
            doc.add_paragraph(f"Export date: {date.today().isoformat()}")
            doc.add_paragraph('')

            for i, item in enumerate(data, start=1):
                title = item.get('title', 'Untitled')
                page_date = item.get('date', 'No date')
                assignee = item.get('assignee', 'Unassigned')
                content = item.get('content', '')

                # Title and metadata
                doc.add_heading(f"{i}. {title}", level=2)
                doc.add_paragraph(f"Date: {page_date}")
                doc.add_paragraph(f"Assignee: {assignee}")
                doc.add_paragraph(f"URL: {item.get('url', '')}")
                doc.add_paragraph('')

                # Content — preserve paragraphs split by newline
                if content:
                    for para in content.splitlines():
                        # skip empty lines to avoid many empty paragraphs
                        if para.strip():
                            doc.add_paragraph(para)
                        else:
                            doc.add_paragraph('')

                # Add page break between entries (but not after last)
                if i != len(data):
                    doc.add_page_break()

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'notion_export_{date.today().isoformat()}.docx'
            )

        else:
            return jsonify({"error": f"Unsupported format: {format}"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
